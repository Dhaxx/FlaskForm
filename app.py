from flask import Flask, render_template, request, make_response, send_file
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
import datetime

def format_cpf(cpf):
    cpf = cpf.zfill(11)  # Preenche com zeros à esquerda até ter 11 dígitos
    return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('form.html')

@app.route('/generate_document', methods=['POST'])
def generate_document():
    # Pega os dados do formulário
    subject = request.form.get('subject')
    colaborador = request.form.get('colaborador')
    descricao = request.form.get("descricao")
    entidade = request.form.get("entidade")

    # Cria o documento e escreve o mesmo
    document = Document()
    header = document.sections[0].header
    htable = header.add_table(1, 1, width=Inches(8))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture('./static/logo.jpg', width=Inches(6), height=Inches(2.0))

    # Adiciona o título de seção centralizado com tamanho de fonte 20
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run("ATESTADO DE VISITA TÉCNICA")
    title_run.bold = True
    title_run.font.size = Pt(20)

    data_atual = datetime.datetime.now()
    data_formatada = data_atual.strftime("%d/%m/%Y") 

    # Adiciona a seção de informações
    info_paragraph = document.add_paragraph()
    info_paragraph.add_run(f"A entidade {entidade.title()}, ")
    info_paragraph.add_run("atesta para os devidos fins, que a ")
    info_paragraph.add_run("Empresa Amendola & Amendola Software Ltda").bold = True
    info_paragraph.add_run(f", inscrita no CNPJ nº 04.326.049/0001-90, realizou visita técnica nesta entidade na data de  ")
    info_paragraph.add_run(f"{data_formatada}, ").bold = True
    info_paragraph.add_run("conforme informações abaixo:")

    # Insere os campos dinâmicos no documento
    fields = []
    for key, value in request.form.items():
        if key.startswith('name'):
            index = key.replace('name', '')
            cpf = request.form.get('cpf' + index)
            field = {'name': value, 'area': request.form.get('area' + index), 'cpf': format_cpf(cpf), 'cargo': request.form.get('cargo' + index)}
            fields.append(field)
    
    p_colaborador = document.add_paragraph()
    p_colaborador.add_run(f'Colaborador: ').bold = True
    p_colaborador.add_run(colaborador)

    p_assunto = document.add_paragraph()
    p_assunto.add_run('Assunto: ').bold = True
    p_assunto.add_run(subject)
    # document.add_paragraph() 

    document.add_heading('Descrição do Atendimento Prestado', level=1)
    document.add_paragraph(descricao)

    # Insere os campos dinâmicos no documento
    document.add_heading('Servidores', level=1)
    for field in fields:
        p_area = document.add_paragraph()
        p_area.add_run('Nome: ').bold = True
        p_area.add_run(field['name'])

        p_area = document.add_paragraph()
        p_area.add_run('cpf: ').bold = True
        p_area.add_run(field['cpf'])

        p_area = document.add_paragraph()
        p_area.add_run('Setor: ').bold = True
        p_area.add_run(field['area'])

        p_area = document.add_paragraph()
        p_area.add_run('Cargo: ').bold = True
        p_area.add_run(field['cargo'])
        

        p_assinatura = document.add_paragraph()
        p_assinatura.add_run("Assinatura: ").bold = True
        p_assinatura.add_run("________________________________")
        document.add_paragraph()  # Adicione uma nova linha em branco entre cada conjunto de campos

    # Adiciona os campos de assinatura centralizados
    # document.add_paragraph()
    # p_assinatura_prefeitura = document.add_paragraph()
    # p_assinatura_prefeitura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p_assinatura_prefeitura.add_run("________________________________")
    # p_assinatura_prefeitura.add_run(f"\n{entidade.upper()}").bold = True
    # p_assinatura_prefeitura.add_run("\n(NOME SERVIDOR RESPONSAVEL PELO SETOR)")
    # p_assinatura_prefeitura.add_run("\n(FUNÇÃO SERVIDOR)")

    p_assinatura_empresa = document.add_paragraph()
    p_assinatura_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_assinatura_empresa.add_run("________________________________")
    p_assinatura_empresa.add_run("\nAMENDOLA & AMENDOLA SOFTWARE LTDA.").bold = True
    p_assinatura_empresa.add_run(f"\n{colaborador.upper()}")
    p_assinatura_empresa.add_run("\nTÉCNICO RESPONSÁVEL")

    # Salva o documento em um arquivo temporário
    temp_file = 'temp.docx'
    document.save(temp_file)


    # Baixa no navegador
    response = make_response(send_file(temp_file, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
    response.headers.set('Content-Disposition', 'attachment', filename='document.docx')
    return response

    # Abre o documento
    # subprocess.Popen(['start', temp_file], shell=True)

    # return 'Documento gerado com sucesso!'
if __name__ == '__main__':
    app.run(host="0.0.0.0")